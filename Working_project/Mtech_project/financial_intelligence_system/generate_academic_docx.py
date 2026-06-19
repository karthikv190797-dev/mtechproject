"""
Convert ACADEMIC_DOCUMENTATION.md into a formatted Word (.docx) document.

Usage:
    python generate_academic_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


HERE = Path(__file__).parent
SOURCE_MD = HERE / "ACADEMIC_DOCUMENTATION.md"
OUTPUT_DOCX = HERE / "ACADEMIC_DOCUMENTATION.docx"

HEADING_COLOR = RGBColor(0x0B, 0x33, 0x66)
CODE_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
LINK_COLOR = RGBColor(0x00, 0x00, 0xCC)

INLINE_PATTERN = re.compile(r"(\*\*.*?\*\*|`[^`]*`|\[[^\]]*\]\([^)]*\))")


def add_inline(paragraph, text):
    """Add text to a paragraph honouring **bold**, `code`, and [link](url)."""
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_COLOR
        elif part.startswith("[") and "](" in part:
            match = re.match(r"\[(.*?)\]\((.*?)\)", part)
            if match:
                run = paragraph.add_run(match.group(1))
                run.underline = True
                run.font.color.rgb = LINK_COLOR
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def add_code_block(doc, code_lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = CODE_COLOR


def add_table(doc, table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(num_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            text = row[c_idx] if c_idx < len(row) else ""
            text = text.replace("**", "").replace("`", "")
            para = cell.paragraphs[0]
            run = para.add_run(text)
            if r_idx == 0:
                run.bold = True
    doc.add_paragraph()


def is_table_row(line):
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_separator(line):
    return bool(re.match(r"^\|?[\s:|-]+\|?$", line.strip())) and "-" in line


def convert(md_text, doc):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", stripped):
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            heading = doc.add_heading(text, level=min(level, 4))
            for run in heading.runs:
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        # Code block (fenced)
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip closing fence
            continue

        # Tables
        if is_table_row(line):
            table_lines = []
            while i < n and is_table_row(lines[i]):
                if not is_separator(lines[i]):
                    table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            add_inline(para, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            run_para = para.add_run()
            add_inline(para, text)
            for run in para.runs:
                run.italic = True
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        add_inline(para, stripped)
        i += 1


def build_title_page(doc):
    title = doc.add_heading("Self-Correcting Financial Intelligence System", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph(
        "A Multi-Agent Retrieval-Augmented Generation Pipeline with "
        "Teacher-Student Distillation and Evaluator-in-the-Loop Self-Correction"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph("M.Tech Dissertation / Technical Report")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(12)
    meta.runs[0].bold = True

    ver = doc.add_paragraph("System Version 1.0.0  |  Generated June 17, 2026")
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def main():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")

    md_text = SOURCE_MD.read_text(encoding="utf-8")

    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    build_title_page(doc)
    convert(md_text, doc)

    doc.save(OUTPUT_DOCX)
    print(f"Word document created: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
