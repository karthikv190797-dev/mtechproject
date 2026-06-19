from pathlib import Path
import json
import re
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_metadata(pdf_text: str):
    def pick(pattern: str, default: str):
        m = re.search(pattern, pdf_text)
        return m.group(1).strip() if m else default

    return {
        "student_name": pick(r"NAME OF THE STUDENT\s+([^\n]+)", "Karthik V"),
        "student_id": pick(r"STUDENT ID No\.\s*([^\n]+)", "2024AA05045"),
        "course_no": pick(r"Course\s*No\.:\s*([^\n]+)", "AIMLCZG628T"),
        "degree": pick(r"Degree\s*Program:\s*([^\n]+)", "M.Tech in Artificial Intelligence and Machine Learning"),
        "organization": pick(r"Organization:\s*([^\n]+)", "Voya Global Services Pvt Ltd"),
        "location": pick(r"Location:\s*([^\n]+)", "Bangalore, India"),
        "supervisor": pick(r"SUPERVISOR'S NAME\s+([^\n]+)", "ShivaChandran Masilamani"),
        "examiner": pick(r"ADDITIONAL EXAMINER'S\s+NAME\s+([^\n]+)", "Deepthi B"),
    }


def section_after(text: str, marker: str, limit: int = 3000) -> str:
    idx = text.find(marker)
    if idx < 0:
        return ""
    part = text[idx: idx + limit]
    part = re.sub(r"--- PAGE \d+ ---", "", part)
    part = re.sub(r"\n{3,}", "\n\n", part)
    return part.strip()


def add_cover_page(doc: Document, meta: dict):
    p = doc.add_paragraph("BIRLA INSTITUTE OF TECHNOLOGY AND SCIENCE, PILANI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("WORK INTEGRATED LEARNING PROGRAMMES (WILP)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")
    p = doc.add_paragraph("FINAL PROJECT REPORT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    p = doc.add_paragraph("Self-Correcting Financial Intelligence System")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(15)

    doc.add_paragraph("")
    doc.add_paragraph(f"Student Name: {meta['student_name']}")
    doc.add_paragraph(f"BITS ID: {meta['student_id']}")
    doc.add_paragraph(f"Course No.: {meta['course_no']}")
    doc.add_paragraph(f"Degree Programme: {meta['degree']}")
    doc.add_paragraph(f"Organization: {meta['organization']}")
    doc.add_paragraph(f"Location: {meta['location']}")
    doc.add_paragraph(f"Supervisor: {meta['supervisor']}")
    doc.add_paragraph(f"Additional Examiner: {meta['examiner']}")
    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")


def main():
    root = Path(r"e:\Final_Mtech_Project\Mtech_project")
    extracted_path = root / "Documents" / "extracted" / "abstract_sources_extracted.json"

    if not extracted_path.exists():
        raise FileNotFoundError(f"Missing extracted source file: {extracted_path}")

    src = json.loads(extracted_path.read_text(encoding="utf-8"))
    pdf_text = src.get("pdf_text", "")
    _pptx_text = src.get("pptx_text", "")

    meta = parse_metadata(pdf_text)
    background_src = section_after(pdf_text, "2.  Background", 5000)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    add_cover_page(doc, meta)

    doc.add_page_break()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This project presents a Self-Correcting Financial Intelligence System that combines a multi-agent "
        "Retrieval-Augmented Generation (RAG) pipeline with evaluator-driven feedback for trustworthy financial analytics. "
        "The system integrates hybrid retrieval using FAISS and BM25 over structured and unstructured financial sources, "
        "including SEC filings and enterprise datasets. A teacher-student distillation strategy is used to balance deep "
        "reasoning quality and runtime efficiency.\n\n"
        "The workflow is orchestrated through specialized agents for retrieval, reading, reasoning, evaluation, verification, "
        "and correction. Quality control is enforced using relevance, factual consistency, numerical coherence, and entailment "
        "scores. Responses below confidence thresholds trigger automatic correction cycles and cross-model verification. "
        "The platform is implemented as a modular, containerized solution using FastAPI, Streamlit, Docker, and YAML-based "
        "configuration to support secure on-premise enterprise deployment.\n\n"
        "Results from implementation and validation indicate improved reliability, reduced hallucinations, and better "
        "traceability of generated financial insights compared to static single-pass RAG approaches."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Financial analysis workflows increasingly depend on extracting accurate insight from lengthy, heterogeneous, and "
        "rapidly growing sources such as annual reports, SEC 10-K filings, earnings transcripts, and data warehouses. "
        "Traditional manual analysis and generic AI pipelines are often slow, non-repeatable, and vulnerable to factual "
        "inconsistencies. This project addresses the gap by building an enterprise-grade, self-correcting financial "
        "intelligence system with transparent and auditable AI reasoning."
    )

    doc.add_heading("2. Background and Motivation", level=1)
    if background_src:
        doc.add_paragraph(re.sub(r"\s+", " ", background_src[:3500]).strip())
    else:
        doc.add_paragraph(
            "The project is motivated by limitations in static RAG systems: insufficient verification, privacy risks in "
            "cloud-only deployment, weak numerical grounding, and inability to self-correct. The proposed solution introduces "
            "evaluator-in-the-loop control and cross-model verification for trustworthy enterprise usage."
        )

    doc.add_heading("3. Objectives", level=1)
    objectives = [
        "Design a LangGraph-based multi-agent RAG framework for retrieval, reasoning, evaluation, verification, and correction.",
        "Implement hybrid retrieval combining dense semantic FAISS search with sparse BM25 keyword matching.",
        "Build a teacher-student reasoning pipeline to optimize quality and inference cost.",
        "Integrate evaluator-in-the-loop self-correction using confidence-based triggering.",
        "Develop cross-model verification to reduce hallucination and model bias.",
        "Support ingestion from both structured and unstructured financial data sources.",
        "Provide API-first backend and interactive UI for analyst workflows.",
        "Enable secure, containerized, on-premise deployment with configuration-driven operation.",
        "Measure retrieval and response quality with standard IR and faithfulness metrics.",
    ]
    for item in objectives:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("4. Scope of Work", level=1)
    doc.add_paragraph(
        "The scope includes end-to-end architecture design, implementation, and validation of the system across data ingestion, "
        "retrieval, reasoning, evaluation, orchestration, backend APIs, and frontend interaction. The implementation is limited "
        "to financial document and table understanding use cases, with emphasis on trustworthiness and enterprise deployability. "
        "The work excludes full-scale production hardening for internet-facing deployment and non-financial vertical adaptation."
    )

    doc.add_heading("5. Methodology and System Design", level=1)
    sections = [
        ("5.1 Data Ingestion Layer", "Multi-source ingestion from SEC filings, Snowflake, and local documents with normalized document representation and metadata pipeline."),
        ("5.2 Hybrid Retrieval Engine", "Dual retrieval with FAISS for semantic relevance and BM25 for lexical precision. Final ranking uses weighted score fusion and top-k evidence selection."),
        ("5.3 Reasoning Pipeline", "Teacher-student framework where a high-capacity model guides a smaller model for efficient financial Q&A and numerical reasoning."),
        ("5.4 Evaluator Agent", "Automated quality assessment across relevance, factual consistency, numerical coherence, and entailment."),
        ("5.5 Correction and Verification", "Low-confidence outputs trigger correction loops and verifier checks for consensus and grounding."),
        ("5.6 Orchestration Layer", "LangGraph-controlled stateful workflow with dynamic routing and iteration limits."),
        ("5.7 Interfaces", "FastAPI backend for programmatic use and Streamlit frontend for analyst interaction, inspection, and traceability."),
    ]
    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    doc.add_heading("6. Implementation Summary", level=1)
    impl_points = [
        "Core modules completed: ingestion, retrieval, reasoning, evaluation, orchestration, backend API, frontend UI, and configuration.",
        "Hybrid retriever supports scalable document retrieval with ranked evidence outputs.",
        "Multi-agent workflow includes Retriever, Reader, Reasoning, Evaluator, Verifier, and Correction agents.",
        "Backend endpoints include /health, /query, /query/{id}, /batch, and /statistics.",
        "Containerization completed with Dockerfile, Dockerfile.frontend, and docker-compose.",
    ]
    for item in impl_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Evaluation and Validation Framework", level=1)
    doc.add_paragraph("The system is evaluated using both retrieval and generation quality metrics. Planned metrics include:")
    metrics = [
        "Precision@K",
        "Mean Reciprocal Rank (MRR)",
        "NDCG",
        "Faithfulness Score",
        "Numerical Consistency",
        "Hallucination Detection Accuracy",
        "Cross-Model Agreement Rate",
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style="List Bullet")

    doc.add_paragraph(
        "Ablation studies compare baseline RAG versus variants with evaluator loop, hybrid retrieval, and verification enabled."
    )

    doc.add_heading("8. Progress from Mid-Sem Abstract to Final Deliverable", level=1)
    doc.add_paragraph(
        "Compared to the mid-sem proposal stage, the final project now includes full implementation of the multi-agent orchestration, "
        "hybrid retrieval stack, evaluator-driven self-correction logic, and user-facing/API interfaces. Deployment artifacts and "
        "configuration-driven execution are complete, making the system demonstrable as an end-to-end enterprise prototype."
    )

    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        "The Self-Correcting Financial Intelligence System demonstrates that evaluator-in-the-loop and cross-model validation can "
        "significantly improve trustworthiness in financial AI assistants. By combining modular architecture, hybrid retrieval, "
        "and containerized deployment, the project bridges research concepts with practical enterprise applicability."
    )

    doc.add_heading("10. Future Work", level=1)
    future_items = [
        "Domain adaptation on larger proprietary financial corpora.",
        "Automated benchmark suite and continuous evaluation pipeline.",
        "Enhanced table structure reasoning and chart understanding.",
        "Role-based access control and stronger governance features.",
        "Latency optimization for large-scale concurrent usage.",
    ]
    for item in future_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("References", level=1)
    refs = [
        "Project implementation artifacts and source modules in financial_intelligence_system.",
        "Dissertation abstract material from Final_Project abstract document.pdf.",
        "Presentation material from Self-Correcting_Financial_Intelligence_System.pptx.",
        "RAG, distillation, and evaluation literature cited in project documentation.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    out_path = root / "Final_Project_Report_Self_Correcting_Financial_Intelligence_System.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
