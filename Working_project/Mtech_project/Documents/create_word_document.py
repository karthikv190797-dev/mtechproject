"""
Script to convert markdown documentation to Word document
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def read_markdown_file(filepath):
    """Read markdown file"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return None


def add_markdown_to_docx(doc, markdown_text, filename=""):
    """Convert markdown text to Word document elements"""
    
    if filename and filename != "DOCUMENTATION_INDEX.md":
        # Add filename as section heading
        heading = doc.add_heading(filename.replace('.md', ''), level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph()
    
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at the start
        if not line.strip():
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            heading_text = line.lstrip('# ').strip()
            doc.add_heading(heading_text, level=1)
        elif line.startswith('## '):
            heading_text = line.lstrip('## ').strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith('### '):
            heading_text = line.lstrip('### ').strip()
            doc.add_heading(heading_text, level=3)
        elif line.startswith('#### '):
            heading_text = line.lstrip('#### ').strip()
            doc.add_heading(heading_text, level=4)
        
        # Handle code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                code_para = doc.add_paragraph(code_text, style='Normal')
                code_para_format = code_para.paragraph_format
                code_para_format.left_indent = Inches(0.5)
                
                # Style code
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Handle tables (basic)
        elif ' | ' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = [line]
            i += 1
            
            # Skip separator line
            if '---' in lines[i] or '---' in lines[i]:
                i += 1
            
            # Collect table rows
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1
            
            # Create table
            try:
                rows = [l.split('|')[1:-1] for l in table_lines if '|' in l]
                if rows:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row in enumerate(rows):
                        for col_idx, cell_text in enumerate(row):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text.strip()
            except:
                pass
        
        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = re.sub(r'^[\s]*[-*]\s+', '', line)
            doc.add_paragraph(bullet_text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            list_text = re.sub(r'^\d+\.\s+', '', line)
            doc.add_paragraph(list_text, style='List Number')
        
        # Handle bold text
        elif line.strip():
            # Add as normal paragraph but with formatting
            para = doc.add_paragraph()
            
            # Process inline formatting
            text = line.strip()
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|`.*?`|\[.*?\]\(.*?\))', text)
            
            for part in parts:
                if not part:
                    continue
                
                # Bold
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                # Italic
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                # Code
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                # Link
                elif part.startswith('[') and '](' in part:
                    match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                    if match:
                        text_part = match.group(1)
                        url = match.group(2)
                        run = para.add_run(text_part)
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    else:
                        para.add_run(part)
                else:
                    para.add_run(part)
        
        i += 1


def main():
    """Main function to create Word document"""
    
    # Get project root
    project_root = Path(__file__).parent.parent / 'Mtech_project'
    
    # Create Word document
    doc = Document()
    
    # Add title page
    title = doc.add_heading('Self-Correcting Financial Intelligence System', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.font.size = Pt(28)
    title_format.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_paragraph('M.Tech Dissertation | BITS WILP | June 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Generated on June 17, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Add table of contents heading
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph(
        'This document contains comprehensive documentation for the Financial Intelligence System. '
        'Please use the navigation links or search function to find specific information.'
    )
    
    toc_items = [
        ('DOCUMENTATION.md', 'Main Documentation - Complete System Guide'),
        ('ARCHITECTURE_AND_DESIGN.md', 'Architecture & Design - Technical Deep Dive'),
        ('API_REFERENCE.md', 'REST API Reference - Complete Endpoint Documentation'),
        ('DEPLOYMENT_AND_OPERATIONS.md', 'Deployment & Operations - Deployment Guides'),
        ('EXAMPLES_AND_USAGE.md', 'Examples & Usage - Code Examples and Integration'),
        ('QUICK_REFERENCE.md', 'Quick Reference - Fast Lookup Guide'),
    ]
    
    for idx, (filename, description) in enumerate(toc_items, 1):
        doc.add_paragraph(f"{idx}. {filename.replace('.md', '')}", style='List Number')
        doc.add_paragraph(description, style='List Bullet')
    
    doc.add_page_break()
    
    # Define documentation files to include
    doc_files = [
        'DOCUMENTATION.md',
        'ARCHITECTURE_AND_DESIGN.md',
        'API_REFERENCE.md',
        'DEPLOYMENT_AND_OPERATIONS.md',
        'EXAMPLES_AND_USAGE.md',
        'QUICK_REFERENCE.md',
    ]
    
    # Process each documentation file
    for doc_file in doc_files:
        filepath = project_root / doc_file
        
        if filepath.exists():
            print(f"Processing {doc_file}...")
            markdown_content = read_markdown_file(filepath)
            
            if markdown_content:
                add_markdown_to_docx(doc, markdown_content, doc_file)
                doc.add_page_break()
            else:
                print(f"Warning: Could not read {doc_file}")
        else:
            print(f"Warning: File not found - {filepath}")
    
    # Save document
    output_path = project_root / 'Financial_Intelligence_System_Documentation.docx'
    doc.save(output_path)
    
    print(f"\n✓ Word document created successfully!")
    print(f"✓ File saved to: {output_path}")
    print(f"✓ File size: {os.path.getsize(output_path) / (1024*1024):.2f} MB")


if __name__ == '__main__':
    main()
